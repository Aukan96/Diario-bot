from docx import Document
import json
import os
import shutil
import time

def exportar():
	shutil.copy("conversations.json", "conversations2.json")
	leer = json.loads(open("C:\DBC\conversations2.json").read())

	#obtener informacion del fichero
	today = time.strftime("%d_%m_%y")
	file_status_array = os.stat("C:\DBC\conversations2.json")
	(mode,ino,dev,nlink,uid,gid,size,atime,mtime,ctime) =os.stat("C:\DBC\conversations2.json")
	fecha = time.ctime(ctime)


	#Leer conversacion
	f= leer['_default']
	fr=f['1']
	fra=fr['conversations']
	fra_1=fra[0]
	msgs=[]

	for i in range(len(fra_1)):
	    a=fra_1[i]
	    b=a['msg']
	    msgs.append(b)

	mensajes=' '.join(msgs)

	#creacion de documento word
	document=Document()
	document.add_heading(fecha,0)
	document.add_paragraph(mensajes)
	document.add_page_break()
	document.save(today+'.docx')

	os.remove("C:\DBC\conversations2.json")

	return True